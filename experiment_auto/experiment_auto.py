# sk-lnuxYC2pPFLsK8sxE5KwOBXBDeNiqzckcITqN2xcfV4kHbs2
from docx import Document
import os
import requests
# 新增：从config.py导入配置
from config import KIMI_API_TOKEN, PERSONAL_INFO, INPUT_FOLDER, OUTPUT_FOLDER

# ===================== 配置已移至config.py，无需在此处修改 =====================

# Kimi API基础配置
KIMI_API_URL = "https://api.moonshot.cn/v1/chat/completions"
HEADERS = {
    "Authorization": f"Bearer {KIMI_API_TOKEN}",
    "Content-Type": "application/json; charset=utf-8"
}

def clean_answer(answer):
    """优化清洗：去重+规范代码块换行+去掉多余内容"""
    if not answer:
        return ""
    lines = list(dict.fromkeys([line.strip() for line in answer.splitlines() if line.strip()]))
    cleaned_lines = []
    in_code = False
    for line in lines:
        if line.startswith("```"):
            in_code = not in_code
            cleaned_lines.append(line)
        elif in_code:
            cleaned_lines.append(line)
        else:
            cleaned_lines.append(" ".join(line.split()))
    return "\n".join(cleaned_lines).strip()

def get_kimi_result(full_requirements, core_project):
    """修复：强制4个分点+提高token上限+校验正确性+规范格式"""
    prompt = f"""You are a professional JavaScript teaching assistant. Generate the EXPERIMENT RESULTS in ENGLISH.
STRICT RULES (MUST FOLLOW):
1. EXACTLY 4 points, numbered 1. 2. 3. 4. ONLY, NO EXTRA POINTS.
2. Results MUST be REALISTIC and CORRECT: double-check all calculations and outputs.
3. Keep code snippets SHORT, only show core parts, no redundant comments.
4. Use separate lines for code blocks and list items, no inline code.
5. Focus on the CORE PROJECT: {core_project}
6. No vague words, keep concise and professional.

FULL EXPERIMENT REQUIREMENTS:
{full_requirements}

EXPERIMENT RESULTS (1. 2. 3. 4. ONLY):"""

    payload = {
        "model": "moonshot-v1-8k",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.1,
        "max_tokens": 800
    }
    try:
        response = requests.post(KIMI_API_URL, headers=HEADERS, json=payload, timeout=30)
        response.raise_for_status()
        raw_answer = response.json()["choices"][0]["message"]["content"].strip()
        return clean_answer(raw_answer)
    except Exception as e:
        print(f"⚠️ Result生成失败：{str(e)[:50]}")
        return ""

def get_kimi_conclusion(full_requirements, core_project):
    """修复：严格限制80词内，不超标 + 新增后验字数校验"""
    prompt = f"""You are a professional JavaScript teaching assistant. Write a CONCLUSION in ENGLISH, NO MORE THAN 80 WORDS.
STRICT RULES:
1. Include: 1) What you learned about {core_project}; 2) A small code problem; 3) How you fixed it.
2. NO MORE THAN 80 WORDS, no redundant sentences.
3. Keep professional and natural.

FULL EXPERIMENT REQUIREMENTS:
{full_requirements}

CONCLUSION (80 words MAX):"""

    payload = {
        "model": "moonshot-v1-8k",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "max_tokens": 100
    }
    try:
        response = requests.post(KIMI_API_URL, headers=HEADERS, json=payload, timeout=30)
        response.raise_for_status()
        raw_answer = response.json()["choices"][0]["message"]["content"].strip()
        # 新增：后验字数校验，超过80词自动截断
        word_count = len(raw_answer.split())
        if word_count > 80:
            raw_answer = " ".join(raw_answer.split()[:80])
            # 确保截断后以句号结尾
            if not raw_answer.endswith((".", "!", "?")):
                raw_answer += "."
        return clean_answer(raw_answer)
    except Exception as e:
        print(f"⚠️ Conclusion生成失败：{str(e)[:50]}")
        return ""

def has_image(paragraph):
    try:
        return bool(paragraph._element.xpath(".//blip"))
    except:
        return False

def fill_personal_info_precise(doc):
    """修复：精准对齐个人信息空格"""
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if "Major：" in para_text and "Grade" in para_text and "Class" in para_text:
            para.text = f"Major：{PERSONAL_INFO['major_class']}                  Grade                             Class                     "
        elif "Name" in para_text and "Student ID" in para_text and "Score" in para_text:
            para.text = f"{PERSONAL_INFO['name']}                               {PERSONAL_INFO['student_id']}                     Score                       "
        if PERSONAL_INFO['name'] in para.text:
            break

def extract_requirements_and_core_project(doc):
    full_requirements = ""
    core_project = ""
    in_requirements = False
    project_keywords = ["Program Concept", "Create the Program File", "Project Files", "Part 1", "Part 2"]
    
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if "2. Requirements" in para_text or "Requirements" in para_text and "2." in para_text:
            in_requirements = True
            continue
        if "3. Result" in para_text or "Result" in para_text and "3." in para_text:
            break
        if in_requirements and para_text:
            full_requirements += para_text + "\n"
            for keyword in project_keywords:
                if keyword in para_text:
                    core_project = para_text.split(keyword)[-1].strip().split("\n")[0]
                    break
    
    if not core_project:
        core_project = "JavaScript basic syntax and operations"
    return full_requirements[:3000], core_project

def locate_and_fill(doc, full_requirements, core_project):
    """修复：只取前4个分点，多余的直接丢弃"""
    print(f"  核心项目：{core_project}")
    print("  正在生成实验Result...")
    result_content = get_kimi_result(full_requirements, core_project)
    print("  正在生成实验Conclusion...")
    conclusion_content = get_kimi_conclusion(full_requirements, core_project)
    
    result_points = {}
    if result_content:
        lines = result_content.splitlines()
        current_num = None
        current_text = ""
        for line in lines:
            line = line.strip()
            if line.startswith(("1.", "2.", "3.", "4.")):
                if current_num:
                    result_points[current_num] = current_text.strip()
                current_num = line[0]
                current_text = line[2:].strip()
            elif current_num:
                current_text += "\n" + line
        if current_num and int(current_num) <=4:
            result_points[current_num] = current_text.strip()
    
    in_result = False
    in_conclusion = False
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if "3. Result" in para_text or "Result" in para_text and "3." in para_text:
            in_result = True
            continue
        if "4. Conclusion" in para_text or "Conclusion" in para_text and "4." in para_text:
            in_result = False
            in_conclusion = True
            continue
        if not para_text or has_image(para):
            continue
        
        if in_result:
            if para_text.startswith(("(1)", "(2)", "(3)", "(4)")):
                num = para_text[1]
                if num in result_points:
                    para.clear()
                    para.add_run(f"({num})")
                    para.add_run(f"\n\n{result_points[num]}")
        
        if in_conclusion and conclusion_content:
            if "Write your conclusion" in para_text or not para_text:
                para.clear()
                para.add_run(conclusion_content)
                in_conclusion = False

def process_single_doc(input_path, output_path):
    doc = Document(input_path)
    file_name = os.path.basename(input_path)
    print(f"正在处理：{file_name}")
    
    fill_personal_info_precise(doc)
    full_requirements, core_project = extract_requirements_and_core_project(doc)
    if not full_requirements:
        print(f"⚠️ 未提取到实验要求，跳过：{file_name}")
        doc.save(output_path)
        return
    locate_and_fill(doc, full_requirements, core_project)
    
    doc.save(output_path)
    print(f"✅ 处理完成：{file_name}\n")

def batch_process_all():
    os.makedirs(INPUT_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    doc_files = [f for f in os.listdir(INPUT_FOLDER) if f.endswith(".docx")]
    
    if not doc_files:
        print(f"⚠️ 【{INPUT_FOLDER}】文件夹中未找到任何.docx格式的实验报告，请检查！")
        return
    
    print(f"📂 检测到{len(doc_files)}个待处理实验报告，开始批量处理...\n" + "-"*80 + "\n")
    for file_name in doc_files:
        input_path = os.path.join(INPUT_FOLDER, file_name)
        output_path = os.path.join(OUTPUT_FOLDER, f"finished_{file_name}")
        process_single_doc(input_path, output_path)
    
    print("-"*80 + f"\n🎉 所有实验报告处理完成！最终版已保存至【{OUTPUT_FOLDER}】文件夹")

if __name__ == "__main__":
    batch_process_all()