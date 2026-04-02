from docx import Document
import os
import requests

# 配置导入：优先从项目根目录导入，失败时尝试添加路径
try:
    from config import KIMI_API_TOKEN, PERSONAL_INFO, INPUT_FOLDER, OUTPUT_FOLDER
except ImportError:
    import sys
    # 添加项目根目录到路径（支持从子目录运行）
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    if project_root not in sys.path:
        sys.path.insert(0, project_root)
    from config import KIMI_API_TOKEN, PERSONAL_INFO, INPUT_FOLDER, OUTPUT_FOLDER

# ===================== 配置已移至config.py，无需在此处修改 =====================

# Kimi API基础配置
KIMI_API_URL = "https://api.moonshot.cn/v1/chat/completions"
HEADERS = {
    "Authorization": f"Bearer {KIMI_API_TOKEN}",
    "Content-Type": "application/json; charset=utf-8"
}

def clean_answer(answer):
    """清洗答案：去空行/首尾空格/重复内容，保证1-3行简洁"""
    if not answer:
        return ""
    lines = list(dict.fromkeys([line.strip() for line in answer.splitlines() if line.strip()]))
    cleaned = "\n".join(lines[:3]).strip()
    return cleaned

def get_kimi_answer(question):
    """重写Prompt：强制精准作答，杜绝敷衍/瞎写/重复，图片题直接返回空"""
    is_css_selector = "CSS selectors" in question and "Hello World" in question
    is_display = "display property" in question.lower()
    is_image = "image" in question.lower() or "create a WEB page looks like the following" in question

    prompt_base = """You are a professional front-end tutor, answer in ENGLISH, 1-3 lines ONLY, NO REPEAT, NO NONSENSE.
STRICT RULES (MUST FOLLOW):
1. Answer must be SPECIFIC and ACCURATE, NO vague words like "several/some/many".
2. NO duplicate content, one answer only for one question.
3. If you can't answer with pure text, return EMPTY STRING directly, no any words.
4. Keep professional, fit for college assignments, no redundant words."""

    if is_image:
        return ""
    elif is_css_selector:
        prompt_base += """
FOR THIS CSS SELECTOR QUESTION:
1. Use ONLY STANDARD CSS selectors (class/descendant/child/attribute/element), NO :contains().
2. Must match the HTML structure: <div class="container"><p class="message">Hello World</p></div>.
3. List AT LEAST 5 valid selectors, concise, 1-3 lines, no comments."""
    elif is_display:
        prompt_base += """
FOR THIS DISPLAY PROPERTY QUESTION:
1. List ALL COMMON useful values directly, NO vague words, 1 line only.
2. Values must include: block, inline, inline-block, none, flex, grid."""

    prompt = f"{prompt_base}\n\nQuestion: {question}\nAnswer:"

    payload = {
        "model": "moonshot-v1-8k",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.1,
        "max_tokens": 150
    }
    try:
        response = requests.post(KIMI_API_URL, headers=HEADERS, json=payload, timeout=25)
        response.raise_for_status()
        raw_answer = response.json()["choices"][0]["message"]["content"].strip()
        return clean_answer(raw_answer)
    except Exception as e:
        print(f"⚠️ 题目处理失败：{question[:30]}... | 原因：{str(e)[:50]}")
        return ""

def has_image(paragraph):
    """检测图片，兼容所有python-docx版本"""
    try:
        return bool(paragraph._element.xpath(".//blip"))
    except:
        return False

def is_question(pure_text):
    """识别题目，覆盖常见题号格式"""
    if len(pure_text) < 5:
        return False
    if pure_text[0].isdigit():
        return True
    if pure_text.startswith(("(1)", "(2)", "（1）", "（2）", "[1]", "[2]")):
        return True
    first_word = pure_text.split()[0].lower()
    if first_word in ["what", "how", "why", "explain", "define", "describe", "list", "try"]:
        return True
    return False

def is_invalid_answer(answer, question):
    """精准校验无效答案：敷衍/空/纯数字/无关键词，直接空白"""
    if not answer or len(answer) < 3:
        return True
    if answer.isdigit() or (answer.replace(".", "", 1).isdigit() and answer.count(".") == 1):
        return True
    if "display property" in question.lower() and not any(v in answer for v in ["block", "inline", "none"]):
        return True
    if "CSS selectors" in question and not any(v in answer for v in [".container", ".message", "p", ">", "[]"]):
        return True
    return False

def fill_personal_info_precise(doc):
    """精准填充个人信息，对齐空格格式"""
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if "Date：" in para_text and "Major" in para_text and "Class" in para_text:
            para.text = f"Date：{PERSONAL_INFO['date']}                    {PERSONAL_INFO['major_class']}                   "
        elif "Name：" in para_text and "StudentID：" in para_text and "Score：" in para_text:
            para.text = f"Name：{PERSONAL_INFO['name']}                  StudentID：{PERSONAL_INFO['student_id']}                   Score：                  "
        if PERSONAL_INFO['name'] in para.text:
            break

def process_single_doc(input_path, output_path):
    """主处理：精准填信息+一题一答+无敷衍+无标记+空白跳过"""
    doc = Document(input_path)
    fill_personal_info_precise(doc)
    processed_questions = []
    processed_para_ids = []
    question_count = 0

    for para in doc.paragraphs:
        para_id = id(para)
        if not para.text.strip() or has_image(para) or para_id in processed_para_ids:
            continue
        
        original_text = para.text.strip()
        pure_question = original_text.splitlines()[0].strip()
        processed_para_ids.append(para_id)
        
        if not is_question(pure_question):
            continue
        
        if pure_question in processed_questions:
            continue
        
        question_count += 1
        print(f"  [处理题目 {question_count}] {pure_question[:60]}...")
        para.clear()
        para.add_run(pure_question)
        processed_questions.append(pure_question)
        
        answer = get_kimi_answer(pure_question)
        
        if not is_invalid_answer(answer, pure_question):
            para.add_run(f"\n\n{answer}")

    doc.save(output_path)
    print(f"✅ 处理完成 | 有效题目：{question_count}道 | 保存：{os.path.basename(output_path)}\n")

def batch_process_all():
    """批量处理所有文档，自动创建文件夹"""
    os.makedirs(INPUT_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    doc_files = [f for f in os.listdir(INPUT_FOLDER) if f.endswith(".docx")]
    
    if not doc_files:
        print(f"⚠️ 【{INPUT_FOLDER}】文件夹中未找到任何.docx格式的文档，请检查！")
        return
    
    print(f"📂 检测到{len(doc_files)}个待处理作业，开始批量处理（已填充个人信息）...\n" + "-"*80 + "\n")
    for file_name in doc_files:
        print(f"正在处理：{file_name}")
        input_path = os.path.join(INPUT_FOLDER, file_name)
        output_path = os.path.join(OUTPUT_FOLDER, f"finished_{file_name}")
        process_single_doc(input_path, output_path)
    
    print("-"*80 + f"\n🎉 所有作业处理完成！最终版文档已保存至【{OUTPUT_FOLDER}】文件夹")

if __name__ == "__main__":
    batch_process_all()
